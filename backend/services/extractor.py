"""
Stage 1 — Document Ingestion

Extraction pipeline (in priority order):
  PDF  → pdfplumber → PyMuPDF → pytesseract OCR (scanned)
  DOCX → python-docx (paragraphs + tables)
  Image → pytesseract OCR with OpenCV preprocessing
  TXT  → plain read

Returns cleaned, normalised plain text ready for segmentation.
"""

import os
import re
import io
from pathlib import Path
from typing import Optional


# ── Public entry point ────────────────────────────────────────────────────────

def extract_text(file_path: str, ext: Optional[str] = None) -> str:
    """Return cleaned text from any supported document format."""
    if ext is None:
        ext = Path(file_path).suffix
    ext = ext.lower()

    if ext == ".txt":
        return _extract_txt(file_path)
    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext in {".docx", ".doc"}:
        return _extract_docx(file_path)
    if ext in {".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp"}:
        return _extract_image_ocr(file_path)
    raise ValueError(f"Unsupported file extension: {ext}")


# ── TXT ───────────────────────────────────────────────────────────────────────

def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return _clean(f.read())


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    """Three-stage PDF extraction: pdfplumber → PyMuPDF → OCR."""

    # Stage A: pdfplumber — best for digitally created PDFs
    text = _pdfplumber_extract(path)
    if _text_is_usable(text):
        return _clean(text)

    # Stage B: PyMuPDF — handles more complex PDF structures
    text = _pymupdf_extract(path)
    if _text_is_usable(text):
        return _clean(text)

    # Stage C: OCR — for scanned / image-only PDFs
    print(f"[extractor] PDF appears scanned — running OCR on {Path(path).name}")
    return _clean(_pdf_ocr(path))


def _pdfplumber_extract(path: str) -> str:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=3, y_tolerance=3)
                if t:
                    pages.append(t)
        return "\n\n".join(pages)
    except Exception as e:
        print(f"[extractor] pdfplumber failed: {e}")
        return ""


def _pymupdf_extract(path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n\n".join(p for p in pages if p.strip())
    except Exception as e:
        print(f"[extractor] PyMuPDF failed: {e}")
        return ""


def _pdf_ocr(path: str) -> str:
    """Render each PDF page as an image then run Tesseract OCR."""
    try:
        import fitz
        doc = fitz.open(path)
        results = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            results.append(_ocr_bytes(img_bytes))
        doc.close()
        return "\n\n".join(r for r in results if r.strip())
    except Exception as e:
        print(f"[extractor] PDF OCR failed: {e}")
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Extract text from tables too — clauses are often hidden there
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return _clean("\n\n".join(parts))
    except Exception as e:
        print(f"[extractor] DOCX extraction failed: {e}")
        return ""


# ── Image OCR ─────────────────────────────────────────────────────────────────

def _extract_image_ocr(path: str) -> str:
    with open(path, "rb") as f:
        return _clean(_ocr_bytes(f.read()))


def _ocr_bytes(img_bytes: bytes) -> str:
    """Run Tesseract OCR with OpenCV preprocessing on raw image bytes."""
    try:
        import pytesseract
        import numpy as np
        import cv2
        from PIL import Image

        # Decode image
        nparr = np.frombuffer(img_bytes, np.uint8)
        img   = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img is None:
            img = np.array(Image.open(io.BytesIO(img_bytes)).convert("RGB"))

        # Preprocessing: greyscale → denoise → binarise → deskew
        grey  = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        grey  = cv2.fastNlMeansDenoising(grey, h=10)
        _, bw = cv2.threshold(grey, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Deskew
        coords = np.column_stack(np.where(bw < 128))
        if len(coords) > 100:
            angle = cv2.minAreaRect(coords)[-1]
            if abs(angle) < 45:
                h, w = bw.shape
                M   = cv2.getRotationMatrix2D((w // 2, h // 2), angle if angle > -45 else angle + 90, 1.0)
                bw  = cv2.warpAffine(bw, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

        pil_img = Image.fromarray(bw)
        return pytesseract.image_to_string(pil_img, config="--oem 3 --psm 6", lang="eng")

    except ImportError:
        print("[extractor] pytesseract / opencv not installed — OCR unavailable")
        return ""
    except Exception as e:
        print(f"[extractor] OCR failed: {e}")
        return ""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _text_is_usable(text: str, min_chars: int = 300) -> bool:
    """Return True if extraction produced meaningful content."""
    return bool(text) and len(text.strip()) >= min_chars


def _clean(text: str) -> str:
    """Normalise whitespace and remove common PDF artefacts."""
    if not text:
        return ""
    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove standalone page numbers
    text = re.sub(r"(?m)^\s*\d{1,4}\s*$", "", text)
    # Collapse excessive blank lines
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    # Normalise spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
